"""
Comprehensive build readiness tests for CasePrepd Windows installer.

Covers gaps NOT tested by test_installer_paths.py, test_installer_readiness.py,
or test_production_readiness.py:

P0 - Frozen-mode path resolution (Path(__file__).parent.parent fragility)
P1 - Hidden import completeness (spec vs reality)
P1 - DLL-heavy package loading (onnxruntime, torch, tokenizers, etc.)
P1 - Package data accessibility (customtkinter themes, certifi bundle, etc.)
P1 - Download script cross-reference (model dir names match config.py)
P1 - Post-build dist/ verification (skip when no build exists)
P2 - Installer script validation (GUID, version, privileges, architecture)
P2 - MAX_PATH safety (Windows 260-char limit)
P2 - Antivirus awareness (icon, metadata, code signing)
P2 - Upgrade safety (user data isolation from install dir)

Research references:
- PyInstaller hidden import failures: https://pyinstaller.org/en/stable/when-things-go-wrong.html
- Windows MAX_PATH: https://learn.microsoft.com/en-us/windows/win32/fileio/maximum-file-path-limitation
- AV false positives on unsigned EXEs: https://github.com/pyinstaller/pyinstaller/issues/6474
- Inno Setup best practices: https://jrsoftware.org/ishelp/
"""

import importlib
import os
import re
import sys
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).parent.parent
SPEC_PATH = PROJECT_ROOT / "caseprepd.spec"
ISS_PATH = PROJECT_ROOT / "installer" / "caseprepd.iss"
DIST_DIR = PROJECT_ROOT / "dist" / "CasePrepd"


def _read_spec():
    """Read spec file or skip."""
    if not SPEC_PATH.exists():
        pytest.skip("caseprepd.spec not found")
    return SPEC_PATH.read_text(encoding="utf-8")


def _read_iss():
    """Read .iss file or skip."""
    if not ISS_PATH.exists():
        pytest.skip("caseprepd.iss not found")
    return ISS_PATH.read_text(encoding="utf-8")


# ============================================================================
# P0 - Frozen Mode Path Resolution
# ============================================================================
# config.py has 8 constants using Path(__file__).parent.parent instead of
# BUNDLED_BASE_DIR. In PyInstaller onedir, both resolve to _internal/, but
# this is coincidental. These tests catch drift and document the fragility.


class TestFrozenModePathResolution:
    """Tests that Path(__file__)-based constants stay equivalent to BUNDLED_BASE_DIR."""

    def test_dev_mode_paths_equivalent(self):
        """Path(config.__file__).parent.parent == BUNDLED_BASE_DIR in dev mode."""
        import src.config as config

        assert not getattr(sys, "frozen", False), "Test assumes dev mode"
        assert Path(config.__file__).parent.parent == config.BUNDLED_BASE_DIR

    @pytest.mark.parametrize(
        "constant_name",
        [
            "DEFAULT_FEEDBACK_CSV",
            "LEGAL_EXCLUDE_LIST_PATH",
            "MEDICAL_TERMS_LIST_PATH",
            "GOOGLE_WORD_FREQUENCY_FILE",
            "MODEL_CONFIG_FILE",
            "DEBUG_DEFAULT_FILE",
        ],
    )
    def test_all_parent_parent_constants_under_bundled_base(self, constant_name):
        """Each Path(__file__)-based constant resolves under BUNDLED_BASE_DIR."""
        import src.config as config

        value = getattr(config, constant_name)
        try:
            value.relative_to(config.BUNDLED_BASE_DIR)
        except ValueError:
            pytest.fail(
                f"{constant_name} = {value} is not under "
                f"BUNDLED_BASE_DIR = {config.BUNDLED_BASE_DIR}"
            )

    def test_config_constants_use_bundled_config_dir(self):
        """Config-file constants use BUNDLED_CONFIG_DIR, not Path(__file__)."""
        import src.config as config

        source = Path(config.__file__).read_text(encoding="utf-8")

        # These constants point into config/ and must use BUNDLED_CONFIG_DIR
        config_file_constants = [
            "DEFAULT_FEEDBACK_CSV",
            "LEGAL_EXCLUDE_LIST_PATH",
            "MEDICAL_TERMS_LIST_PATH",
            "MODEL_CONFIG_FILE",
        ]
        fragile = []
        for name in config_file_constants:
            pattern = rf"^{name}\s*=.*Path\(__file__\)"
            if re.search(pattern, source, re.MULTILINE):
                fragile.append(name)

        assert not fragile, (
            f"These constants use Path(__file__).parent.parent instead of "
            f"BUNDLED_CONFIG_DIR: {fragile}"
        )

    def test_parent_parent_usage_count_in_source(self):
        """Exactly 1 Path(__file__).parent.parent occurrence (BUNDLED_BASE_DIR definition).

        All other constants should use BUNDLED_BASE_DIR or BUNDLED_CONFIG_DIR.
        If this count increases, a new fragile constant was added.
        """
        import src.config as config

        source = Path(config.__file__).read_text(encoding="utf-8")
        count = source.count("Path(__file__).parent.parent")
        assert count == 1, (
            f"Expected 1 Path(__file__).parent.parent occurrence "
            f"(only the BUNDLED_BASE_DIR definition), found {count}. "
            f"New constants should use BUNDLED_BASE_DIR or BUNDLED_CONFIG_DIR."
        )

    def test_frozen_mode_simulation(self, tmp_path, monkeypatch):
        """Monkeypatch sys.frozen + sys._MEIPASS, reload config, verify BUNDLED_BASE_DIR."""
        fake_meipass = str(tmp_path / "_internal")
        os.makedirs(fake_meipass, exist_ok=True)

        # Create minimal directory structure so config doesn't crash
        config_dir = Path(fake_meipass) / "config"
        config_dir.mkdir(parents=True, exist_ok=True)

        monkeypatch.setattr(sys, "frozen", True, raising=False)
        monkeypatch.setattr(sys, "_MEIPASS", fake_meipass, raising=False)

        # Reload config to pick up frozen mode
        import src.config

        importlib.reload(src.config)
        try:
            assert Path(fake_meipass) == src.config.BUNDLED_BASE_DIR
        finally:
            # Restore dev mode
            monkeypatch.delattr(sys, "frozen", raising=False)
            monkeypatch.delattr(sys, "_MEIPASS", raising=False)
            importlib.reload(src.config)


# ============================================================================
# P1 - Hidden Imports Complete
# ============================================================================
# caseprepd.spec declares 18 hidden imports. Existing tests only verify 8.
# Missing imports = ImportError at runtime in the frozen app.

# All non-UI hidden imports from the spec (tkinterdnd2 tested separately below)
HIDDEN_IMPORTS_NON_UI = [
    "src.core.vocabulary.algorithms.ner_algorithm",
    "src.core.vocabulary.algorithms.rake_algorithm",
    "src.core.vocabulary.algorithms.textrank_algorithm",
    "src.core.vocabulary.algorithms.bm25_algorithm",
    "src.core.vocabulary.algorithms.scispacy_algorithm",
    "src.core.vocabulary.algorithms.yake_algorithm",
    "tiktoken_ext.openai_public",
    "tiktoken_ext",
    "sklearn.utils._typedefs",
    "sklearn.utils._cython_blas",
    "sklearn.neighbors._partition_nodes",
    "sklearn.tree._utils",
]


class TestHiddenImportsComplete:
    """Tests that all spec hidden imports are actually importable."""

    @pytest.mark.parametrize("module_name", HIDDEN_IMPORTS_NON_UI)
    def test_hidden_import_importable(self, module_name):
        """Each non-UI hidden import loads with a valid module structure."""
        mod = importlib.import_module(module_name)
        assert hasattr(mod, "__name__")
        assert mod.__name__ == module_name
        # Must expose attributes (public or private C extension symbols)
        assert len(dir(mod)) > 0, f"{module_name} has no attributes"

    def test_spec_hidden_imports_match_test_list(self):
        """Every hidden import in the spec appears in our test parameter list."""
        spec = _read_spec()

        # Extract the hidden_imports list from spec source
        # Look for lines like:  "module.name",
        pattern = r'"([\w.]+)"'
        # Get just the hidden_imports section
        hi_start = spec.index("hidden_imports = [")
        hi_end = spec.index("]", hi_start)
        hi_section = spec[hi_start:hi_end]

        spec_imports = set(re.findall(pattern, hi_section))
        test_imports = set(HIDDEN_IMPORTS_NON_UI) | {"tkinterdnd2"}

        missing = spec_imports - test_imports
        assert not missing, (
            f"Hidden imports in spec not covered by tests: {missing}. "
            f"Add them to HIDDEN_IMPORTS_NON_UI or the tkinterdnd2 test."
        )


class TestPackagesToCollectImportable:
    """Every package in packages_to_collect must be importable."""

    @staticmethod
    def _extract_packages_to_collect():
        """Parse packages_to_collect list from the spec file."""
        spec = _read_spec()
        start = spec.index("packages_to_collect = [")
        end = spec.index("]", start)
        section = spec[start:end]
        return re.findall(r'"([\w]+)"', section)

    def test_all_packages_importable(self):
        """Each package in packages_to_collect can be imported."""
        import importlib

        packages = self._extract_packages_to_collect()
        assert len(packages) > 0, "No packages found in packages_to_collect"

        failures = []
        for pkg in packages:
            try:
                importlib.import_module(pkg)
            except ImportError as e:
                failures.append(f"{pkg}: {e}")

        assert not failures, "packages_to_collect entries that failed to import:\n" + "\n".join(
            failures
        )

    def test_pytextrank_full_chain(self):
        """pytextrank imports successfully with all dependencies."""
        import pytextrank

        assert hasattr(pytextrank, "TopicRank")
        assert hasattr(pytextrank, "TopicRankFactory")


# ============================================================================
# P1 - DLL-Heavy Package Loading
# ============================================================================
# Spec collects DLLs from torch, tokenizers, thinc, sklearn.
# If DLLs are missing, these packages import but fail on first use.

DLL_HEAVY_PACKAGES = [
    ("torch", "Tensor"),
    ("tokenizers", "Tokenizer"),
    ("thinc.api", "Config"),
    ("sklearn.ensemble", "RandomForestClassifier"),
]


class TestDLLHeavyPackageLoading:
    """Tests that DLL-dependent packages load their native extensions."""

    @pytest.mark.parametrize(
        "package_name,expected_attr",
        DLL_HEAVY_PACKAGES,
        ids=[p[0] for p in DLL_HEAVY_PACKAGES],
    )
    def test_package_loads_with_critical_attribute(self, package_name, expected_attr):
        """Package imports successfully and has its key attribute (proves DLLs loaded)."""
        mod = importlib.import_module(package_name)
        assert hasattr(mod, expected_attr), (
            f"{package_name} imported but missing '{expected_attr}' -- "
            f"native DLLs may not have loaded correctly"
        )


# ============================================================================
# P1 - Package Data Accessibility
# ============================================================================
# 14 packages have data bundled via collect_data_files(). Missing data = crash.


class TestPackageDataAccessibility:
    """Tests that packages with bundled data can access their resources."""

    def test_customtkinter_theme_loadable(self):
        """customtkinter has a themes directory with JSON files."""
        import customtkinter

        pkg_dir = Path(customtkinter.__file__).parent
        theme_dir = pkg_dir / "assets" / "themes"
        if not theme_dir.exists():
            # Some versions use different layout
            theme_dirs = list(pkg_dir.rglob("*.json"))
            assert theme_dirs, (
                "customtkinter has no .json theme files -- "
                "collect_data_files('customtkinter') may be incomplete"
            )
        else:
            json_files = list(theme_dir.glob("*.json"))
            assert json_files, f"No .json files in {theme_dir}"

    def test_certifi_ca_bundle_exists(self):
        """certifi.where() returns an existing CA bundle > 100KB."""
        import certifi

        ca_path = Path(certifi.where())
        assert ca_path.exists(), f"CA bundle not found at {ca_path}"
        size_kb = ca_path.stat().st_size / 1024
        assert size_kb > 100, f"CA bundle suspiciously small: {size_kb:.0f} KB"

    def test_fpdf_functional(self):
        """fpdf creates a working PDF instance with expected methods."""
        from fpdf import FPDF

        pdf = FPDF()
        assert hasattr(pdf, "add_page")
        assert hasattr(pdf, "set_font")
        assert hasattr(pdf, "cell")
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(text="test")
        assert pdf.pages_count == 1

    def test_docx_document_creates(self):
        """python-docx Document() creates a fresh empty document."""
        from docx import Document
        from docx.document import Document as DocumentClass

        doc = Document()
        assert isinstance(doc, DocumentClass)
        assert len(doc.paragraphs) == 0

    def test_nupunkt_sentence_tokenize(self):
        """nupunkt splits two sentences correctly (needs bundled model data)."""
        from nupunkt import sent_tokenize

        result = sent_tokenize("Hello world. How are you?")
        assert len(result) == 2

    def test_spellchecker_initializes(self):
        """SpellChecker loads word frequency data and detects misspellings."""
        from spellchecker import SpellChecker

        spell = SpellChecker()
        unknown = spell.unknown(["asdfqwerty"])
        assert "asdfqwerty" in unknown

    def test_tkinterdnd2_importable(self):
        """tkinterdnd2 imports (TCL scripts + native DLLs for drag-and-drop)."""
        try:
            import tkinterdnd2

            assert tkinterdnd2 is not None
        except (ImportError, RuntimeError) as e:
            if "Tcl" in str(e) or "display" in str(e).lower():
                pytest.skip(f"Tcl/Tk unavailable: {e}")
            raise


# ============================================================================
# P1 - Download Script Cross-Reference
# ============================================================================
# If download_models.py saves to a different name than config.py expects,
# the bundled app fails silently at runtime.


class TestDownloadScriptCrossReference:
    """Tests that model directory names in download script match config.py paths."""

    @pytest.fixture(autouse=True)
    def _import_download_script(self):
        """Import download_models from scripts/."""
        script_path = PROJECT_ROOT / "scripts"
        sys.path.insert(0, str(script_path))
        yield
        sys.path.pop(0)

    @pytest.mark.parametrize(
        "config_attr,expected_suffix",
        [
            ("EMBEDDING_MODEL_LOCAL_PATH", "embeddings/nomic-embed-text-v1.5"),
            ("RERANKER_MODEL_LOCAL_PATH", "gte-reranker-modernbert-base"),
        ],
        ids=[
            "nomic-embed",
            "gte-reranker",
        ],
    )
    def test_hf_model_dirname_matches_config(self, config_attr, expected_suffix):
        """HF model's local_subdir in download script matches config.py path suffix."""
        import download_models

        from src import config

        # Get the config path and verify its suffix matches expected
        config_path = getattr(config, config_attr)
        relative = config_path.relative_to(config.BUNDLED_MODELS_DIR)
        assert str(relative).replace("\\", "/") == expected_suffix

        # Verify the download script has a matching entry
        found = False
        for _, local_subdir, _ in download_models.HF_MODELS:
            if local_subdir == expected_suffix:
                found = True
                break
        assert found, (
            f"No HF_MODELS entry with local_subdir='{expected_suffix}' "
            f"to match config.{config_attr}"
        )

    @pytest.mark.parametrize(
        "config_attr,model_name",
        [
            ("SPACY_EN_CORE_WEB_LG_PATH", "en_core_web_lg"),
            ("SPACY_EN_CORE_WEB_SM_PATH", "en_core_web_sm"),
            ("SPACY_EN_NER_BC5CDR_MD_PATH", "en_ner_bc5cdr_md"),
        ],
    )
    def test_spacy_model_names_match_config(self, config_attr, model_name):
        """spaCy model names in download script match config path constants."""
        import download_models

        from src import config

        config_path = getattr(config, config_attr)
        assert config_path.name == model_name
        assert model_name in download_models.SPACY_MODELS


# ============================================================================
# Nomic Embedding Model — Offline Readiness
# ============================================================================
# The nomic model uses trust_remote_code with custom Python code from a
# separate HF repo (nomic-ai/nomic-bert-2048). These files and a patched
# config.json must be bundled for offline/frozen loading to work.


class TestNomicOfflineReadiness:
    """Verify nomic embedding model is bundled for fully offline use."""

    NOMIC_DIR = PROJECT_ROOT / "models" / "embeddings" / "nomic-embed-text-v1.5"

    def test_custom_config_code_exists(self):
        """configuration_hf_nomic_bert.py is bundled with the model."""
        path = self.NOMIC_DIR / "configuration_hf_nomic_bert.py"
        assert path.is_file(), f"Missing custom config code: {path}"

    def test_custom_modeling_code_exists(self):
        """modeling_hf_nomic_bert.py is bundled with the model."""
        path = self.NOMIC_DIR / "modeling_hf_nomic_bert.py"
        assert path.is_file(), f"Missing custom modeling code: {path}"

    def test_auto_map_uses_local_references(self):
        """config.json auto_map entries point to local modules, not remote repos."""
        import json

        config = json.loads((self.NOMIC_DIR / "config.json").read_text(encoding="utf-8"))
        assert "auto_map" in config, "config.json missing auto_map"
        for key, value in config["auto_map"].items():
            assert "--" not in value, f"auto_map[{key}] still points to remote repo: {value}"

    def test_auto_map_references_bundled_files(self):
        """auto_map module names match the bundled .py files."""
        import json

        config = json.loads((self.NOMIC_DIR / "config.json").read_text(encoding="utf-8"))
        for value in config["auto_map"].values():
            module_name = value.split(".")[0]
            py_file = self.NOMIC_DIR / f"{module_name}.py"
            assert py_file.is_file(), (
                f"auto_map references {module_name} but {py_file.name} not found"
            )

    def test_config_has_nomic_bert_architecture(self):
        """config.json declares NomicBertModel architecture."""
        import json

        config = json.loads((self.NOMIC_DIR / "config.json").read_text(encoding="utf-8"))
        assert "NomicBertModel" in config.get("architectures", [])

    def test_validate_models_requires_custom_code(self):
        """validate_models.py checks for the custom .py files."""
        validate_src = (PROJECT_ROOT / "scripts" / "validate_models.py").read_text(encoding="utf-8")
        assert "configuration_hf_nomic_bert.py" in validate_src
        assert "modeling_hf_nomic_bert.py" in validate_src


# ============================================================================
# P1 - Post-Build Dist Verification
# ============================================================================
# After PyInstaller runs, verify dist/ folder structure before Inno packaging.
# All tests skip when no build exists (CI or dev without build).

DIST_SKIP = pytest.mark.skipif(not DIST_DIR.exists(), reason="No dist/CasePrepd/ build found")


@DIST_SKIP
class TestPostBuildDistVerification:
    """Tests for the PyInstaller dist/ output folder."""

    def test_exe_exists_and_minimum_size(self):
        """CasePrepd.exe exists and is > 1MB."""
        exe = DIST_DIR / "CasePrepd.exe"
        assert exe.is_file(), f"Missing: {exe}"
        size_mb = exe.stat().st_size / (1024 * 1024)
        assert size_mb > 1, f"EXE suspiciously small: {size_mb:.1f} MB"

    def test_internal_directory_exists(self):
        """_internal/ directory exists (onedir mode)."""
        internal = DIST_DIR / "_internal"
        assert internal.is_dir(), f"Missing: {internal}"

    @pytest.mark.parametrize(
        "config_file",
        [
            "config/app_name.txt",
            "config/legal_exclude.txt",
            "config/medical_terms.txt",
            "config/models.yaml",
        ],
    )
    def test_config_files_in_dist(self, config_file):
        """Key config files are present in _internal/."""
        internal = DIST_DIR / "_internal"
        path = internal / config_file
        assert path.is_file(), f"Missing config file in dist: {config_file}"

    def test_customtkinter_themes_in_dist(self):
        """Theme JSON files exist somewhere in _internal/."""
        internal = DIST_DIR / "_internal"
        json_files = list(internal.rglob("customtkinter/**/*.json"))
        assert json_files, "No customtkinter theme JSON files found in dist"

    def test_certifi_bundle_in_dist(self):
        """cacert.pem exists somewhere in _internal/."""
        internal = DIST_DIR / "_internal"
        pem_files = list(internal.rglob("cacert.pem"))
        assert pem_files, "No cacert.pem found in dist _internal/"

    @pytest.mark.parametrize(
        "model_dir",
        [
            "models/spacy",
            "models/tesseract",
            "models/poppler",
            "models/nltk_data",
        ],
    )
    def test_model_directories_in_dist(self, model_dir):
        """Key model directories are present in _internal/."""
        internal = DIST_DIR / "_internal"
        path = internal / model_dir
        assert path.is_dir(), f"Missing model directory in dist: {model_dir}"

    def test_no_pycache_in_dist(self):
        """No __pycache__ directories in dist (bloat)."""
        pycache_dirs = list(DIST_DIR.rglob("__pycache__"))
        assert not pycache_dirs, (
            f"Found {len(pycache_dirs)} __pycache__ dirs in dist: "
            f"{[str(d.relative_to(DIST_DIR)) for d in pycache_dirs[:5]]}"
        )

    def test_no_dev_packages_in_dist(self):
        """No dev-only package top-level directories in dist."""
        internal = DIST_DIR / "_internal"
        # Check for top-level package dirs (not transitive deps that happen to
        # share a prefix, e.g. "mypy_extensions" is a runtime dep of pydantic)
        dev_packages = {
            "pytest": ["pytest"],
            "ruff": ["ruff"],
            "black": ["black"],
        }
        found = []
        for pkg, dir_names in dev_packages.items():
            for name in dir_names:
                if (internal / name).is_dir():
                    found.append(pkg)
        assert not found, f"Dev package directories found in dist: {found}"

    def test_total_size_within_range(self):
        """Dist folder is between 500MB and 10GB."""
        total = sum(f.stat().st_size for f in DIST_DIR.rglob("*") if f.is_file())
        total_mb = total / (1024 * 1024)
        assert 500 < total_mb < 10240, f"Dist size is {total_mb:.0f} MB -- expected 500-10240 MB"


@DIST_SKIP
class TestNomicOfflineReadinessDist:
    """Verify nomic custom code is present in the built dist."""

    NOMIC_DIST = DIST_DIR / "_internal" / "models" / "embeddings" / "nomic-embed-text-v1.5"

    def test_custom_code_in_dist(self):
        """Custom .py files are present in the dist build."""
        assert (self.NOMIC_DIST / "configuration_hf_nomic_bert.py").is_file()
        assert (self.NOMIC_DIST / "modeling_hf_nomic_bert.py").is_file()

    def test_auto_map_local_in_dist(self):
        """Dist config.json auto_map uses local references."""
        import json

        config = json.loads((self.NOMIC_DIST / "config.json").read_text(encoding="utf-8"))
        for key, value in config["auto_map"].items():
            assert "--" not in value, f"Dist auto_map[{key}] still remote: {value}"


# ============================================================================
# P2 - Installer Script Validation
# ============================================================================
# Validates the Inno Setup .iss file for correctness.


class TestInstallerScriptValidation:
    """Tests for installer/caseprepd.iss correctness."""

    def test_appid_is_guid_format(self):
        """AppId is a standard GUID format {XXXXXXXX-XXXX-...}."""
        iss = _read_iss()
        match = re.search(r"AppId=\{\{?([^}]+)\}?", iss)
        assert match, "AppId not found in .iss"
        guid = match.group(1)
        # GUID format: 8-4-4-4-12 hex chars
        assert re.match(
            r"^[0-9A-Fa-f]{8}-[0-9A-Fa-f]{4}-[0-9A-Fa-f]{4}-[0-9A-Fa-f]{4}-[0-9A-Fa-f]{12}$",
            guid,
        ), f"AppId is not a valid GUID: {guid}"

    def test_iss_version_matches_src(self):
        """.iss version matches src/__init__.py __version__."""
        iss = _read_iss()
        from src import __version__

        match = re.search(r'#define\s+MyAppVersion\s+"([^"]+)"', iss)
        assert match, "MyAppVersion not found in .iss"
        iss_version = match.group(1)
        assert iss_version == __version__, (
            f".iss version '{iss_version}' != src.__version__ '{__version__}'"
        )

    def test_privileges_required_lowest(self):
        """PrivilegesRequired=lowest (no admin needed for install)."""
        iss = _read_iss()
        assert "PrivilegesRequired=lowest" in iss

    def test_architectures_x64(self):
        """Architecture is set to x64compatible."""
        iss = _read_iss()
        assert "x64compatible" in iss

    def test_uses_autopf_not_hardcoded(self):
        """{autopf} used for default install dir (not C:\\Program Files)."""
        iss = _read_iss()
        assert "{autopf}" in iss
        # Ensure no hardcoded Program Files path for DefaultDirName
        assert "C:\\Program Files" not in iss

    def test_uninstall_deletes_userappdata(self):
        """Uninstall section cleans up %APPDATA%/CasePrepd."""
        iss = _read_iss()
        assert "{userappdata}" in iss

    def test_vcredist_check_present(self):
        """Installer checks for VC++ redistributable before installing."""
        iss = _read_iss()
        assert "vcruntime140" in iss.lower(), (
            "Installer must check for vcruntime140.dll to prevent cryptic "
            "DLL errors on machines without VC++ redistributable"
        )

    def test_version_info_present(self):
        """Version metadata is set in .iss for Windows file properties."""
        iss = _read_iss()
        assert "VersionInfoVersion" in iss
        assert "VersionInfoCompany" in iss
        assert "VersionInfoDescription" in iss
        assert "VersionInfoProductName" in iss


# ============================================================================
# P2 - MAX_PATH Safety
# ============================================================================
# Windows has a 260-char MAX_PATH limit. Install to C:\Users\VeryLongName\
# AppData\Local\Programs\CasePrepd\ uses ~60 chars as prefix.


class TestMaxPathSafety:
    """Tests that model paths don't exceed Windows MAX_PATH when installed."""

    def test_longest_model_path_under_200_chars(self):
        """Longest relative model path leaves room for 60-char install prefix."""
        models_dir = PROJECT_ROOT / "models"
        if not models_dir.exists():
            pytest.skip("models/ directory not present")

        longest = ""
        for f in models_dir.rglob("*"):
            if f.is_file():
                rel = str(f.relative_to(PROJECT_ROOT)).replace("\\", "/")
                if len(rel) > len(longest):
                    longest = rel

        assert len(longest) < 200, (
            f"Longest model path is {len(longest)} chars: {longest}. "
            f"With a 60-char install prefix, this exceeds MAX_PATH (260)."
        )

    def test_config_path_constants_reasonable_length(self):
        """All Path(__file__)-based config constants are < 100 chars relative."""
        import src.config as config

        constants = [
            "DEFAULT_FEEDBACK_CSV",
            "LEGAL_EXCLUDE_LIST_PATH",
            "MEDICAL_TERMS_LIST_PATH",
            "GOOGLE_WORD_FREQUENCY_FILE",
            "MODEL_CONFIG_FILE",
            "DEBUG_DEFAULT_FILE",
        ]
        for name in constants:
            value = getattr(config, name)
            rel = str(value.relative_to(config.BUNDLED_BASE_DIR)).replace("\\", "/")
            assert len(rel) < 100, f"{name} relative path is {len(rel)} chars: {rel}"

    def test_deepest_model_nesting_under_limit(self):
        """No model path has more than 8 directory levels."""
        models_dir = PROJECT_ROOT / "models"
        if not models_dir.exists():
            pytest.skip("models/ directory not present")

        max_depth = 0
        deepest = ""
        for f in models_dir.rglob("*"):
            if f.is_file():
                rel = f.relative_to(models_dir)
                depth = len(rel.parts)
                if depth > max_depth:
                    max_depth = depth
                    deepest = str(rel)

        assert max_depth <= 8, (
            f"Deepest model path has {max_depth} levels: {deepest}. "
            f"Deep nesting increases MAX_PATH risk."
        )


# ============================================================================
# P2 - Antivirus Awareness
# ============================================================================
# Unsigned PyInstaller EXEs trigger AV false positives. Icon + metadata help.


class TestAntivirusAwareness:
    """Tests for reducing antivirus false positive risk."""

    def test_spec_has_icon(self):
        """Icon is set in spec and the .ico file exists on disk."""
        spec = _read_spec()
        assert "icon=" in spec.lower() or "icon.ico" in spec
        icon_path = PROJECT_ROOT / "assets" / "icon.ico"
        assert icon_path.is_file(), f"Icon file missing: {icon_path}"

    def test_iss_has_version_metadata(self):
        """Version, company, description, product name all present in .iss."""
        iss = _read_iss()
        required_fields = [
            "VersionInfoVersion",
            "VersionInfoCompany",
            "VersionInfoDescription",
            "VersionInfoProductName",
        ]
        for field in required_fields:
            assert field in iss, f"Missing {field} in .iss (helps AV classification)"

    @pytest.mark.xfail(
        reason="No code signing configured. Unsigned EXEs are 10x more likely "
        "to trigger AV false positives. Consider signing with a code "
        "signing certificate (e.g., SignTool + Certum open-source cert).",
        strict=False,
    )
    def test_code_signing_gap_documented(self):
        """Flags that no code signing is configured."""
        spec = _read_spec()
        iss = _read_iss()
        combined = spec + iss
        assert "signtool" in combined.lower() or "codesign" in combined.lower()


# ============================================================================
# P2 - Upgrade Safety
# ============================================================================
# User data (corpus, feedback, preferences) must survive app upgrades.
# All writable paths must be under %APPDATA%, not the install directory.


class TestUpgradeSafety:
    """Tests that user data is isolated from the install directory."""

    def test_appdata_uses_env_variable(self):
        """APPDATA_DIR is derived from %APPDATA% environment variable."""
        from src.config import APPDATA_DIR

        appdata = os.environ.get("APPDATA", "")
        if appdata:
            assert str(APPDATA_DIR).startswith(appdata), (
                f"APPDATA_DIR ({APPDATA_DIR}) doesn't start with %APPDATA% ({appdata})"
            )

    @pytest.mark.parametrize(
        "dir_name",
        [
            "MODELS_DIR",
            "CACHE_DIR",
            "LOGS_DIR",
            "CONFIG_DIR",
            "DATA_DIR",
            "CORPUS_DIR",
            "VECTOR_STORE_DIR",
        ],
    )
    def test_user_data_dirs_under_appdata(self, dir_name):
        """All writable directories are under APPDATA_DIR."""
        from src import config

        dir_path = getattr(config, dir_name)
        try:
            dir_path.relative_to(config.APPDATA_DIR)
        except ValueError:
            pytest.fail(
                f"{dir_name} = {dir_path} is NOT under "
                f"APPDATA_DIR = {config.APPDATA_DIR}. "
                f"User data would be lost on app upgrade."
            )

    def test_no_runtime_writes_to_install_dir(self):
        """APPDATA paths are NOT under BUNDLED_BASE_DIR (install dir)."""
        from src.config import APPDATA_DIR, BUNDLED_BASE_DIR

        try:
            APPDATA_DIR.relative_to(BUNDLED_BASE_DIR)
            pytest.fail(
                f"APPDATA_DIR ({APPDATA_DIR}) is under BUNDLED_BASE_DIR "
                f"({BUNDLED_BASE_DIR}). Writes to install dir fail "
                f"without admin privileges and break on upgrade."
            )
        except ValueError:
            pass  # Good -- APPDATA is NOT under install dir

    def test_installer_uses_autopf(self):
        """.iss uses {autopf} for install directory."""
        iss = _read_iss()
        assert "{autopf}" in iss


# ============================================================================
# P1 - Optional Package Bundling
# ============================================================================
# Two packages (nupunkt, tkinterdnd2) have graceful fallbacks in production code.
# Without proper bundling, end users silently lose features. These tests verify
# both are correctly configured in requirements.txt, caseprepd.spec, and have
# proper fallback + bundled paths.

REQUIREMENTS_PATH = PROJECT_ROOT / "requirements.txt"
DOWNLOAD_SCRIPT_PATH = PROJECT_ROOT / "scripts" / "download_models.py"


def _read_requirements():
    """Read requirements.txt or skip."""
    if not REQUIREMENTS_PATH.exists():
        pytest.skip("requirements.txt not found")
    return REQUIREMENTS_PATH.read_text(encoding="utf-8")


def _read_download_script():
    """Read download_models.py or skip."""
    if not DOWNLOAD_SCRIPT_PATH.exists():
        pytest.skip("scripts/download_models.py not found")
    return DOWNLOAD_SCRIPT_PATH.read_text(encoding="utf-8")


class TestOptionalPackageBundling:
    """Tests that optional packages are properly configured for the installer."""

    # ── 4a. Spec cross-reference ──────────────────────────────────────────

    def test_nupunkt_in_spec_data_packages(self):
        """nupunkt is in spec data_packages (ships bundled model data)."""
        spec = _read_spec()
        data_start = spec.index("data_packages = [")
        data_end = spec.index("]", data_start)
        data_section = spec[data_start:data_end]
        assert '"nupunkt"' in data_section

    def test_tkinterdnd2_in_spec_data_packages(self):
        """tkinterdnd2 is in spec data_packages (ships TCL scripts + DLLs)."""
        spec = _read_spec()
        data_start = spec.index("data_packages = [")
        data_end = spec.index("]", data_start)
        data_section = spec[data_start:data_end]
        assert '"tkinterdnd2"' in data_section

    def test_tkinterdnd2_in_spec_hidden_imports(self):
        """tkinterdnd2 is in spec hidden_imports (lazy-loaded by UI)."""
        spec = _read_spec()
        hi_start = spec.index("hidden_imports = [")
        hi_end = spec.index("]", hi_start)
        hi_section = spec[hi_start:hi_end]
        assert '"tkinterdnd2"' in hi_section

    # ── 4b. Requirements cross-reference ──────────────────────────────────

    def test_nupunkt_in_requirements(self):
        """nupunkt is listed in requirements.txt."""
        reqs = _read_requirements()
        assert re.search(r"^nupunkt", reqs, re.MULTILINE)

    def test_tkinterdnd2_in_requirements(self):
        """tkinterdnd2 is listed in requirements.txt."""
        reqs = _read_requirements()
        assert re.search(r"^tkinterdnd2", reqs, re.MULTILINE)

    # ── 4c. Production fallback audit ─────────────────────────────────────

    def test_nupunkt_has_graceful_fallback(self):
        """sentence_splitter.py has ImportError guard for nupunkt."""
        source_path = PROJECT_ROOT / "src" / "core" / "utils" / "sentence_splitter.py"
        source = source_path.read_text(encoding="utf-8")
        assert "ImportError" in source
        assert "_nupunkt_available" in source or "nupunkt" in source.lower()
        assert "warning" in source.lower() or "logger.warning" in source

    def test_tkinterdnd2_has_graceful_fallback(self):
        """main_window.py has HAS_DND pattern for tkinterdnd2."""
        source_path = PROJECT_ROOT / "src" / "ui" / "main_window.py"
        source = source_path.read_text(encoding="utf-8")
        assert "HAS_DND" in source
        assert "ImportError" in source

    # ── 4d. Model/asset availability ──────────────────────────────────────

    def test_nupunkt_bundled_model_in_package(self):
        """nupunkt package ships its own model data (no separate download)."""
        import nupunkt

        pkg_dir = Path(nupunkt.__file__).parent
        # nupunkt bundles models inside its package directory
        model_files = list(pkg_dir.rglob("*.json.gz")) + list(pkg_dir.rglob("*.bin"))
        assert model_files, (
            f"nupunkt package at {pkg_dir} has no model files -- "
            f"collect_data_files('nupunkt') may be incomplete"
        )

    # ── 4e. Loading pattern validation ────────────────────────────────────

    def test_model_loader_checks_bundled_path(self):
        """model_loader.resolve_model_path checks .exists() before HF fallback."""
        source_path = PROJECT_ROOT / "src" / "core" / "utils" / "model_loader.py"
        source = source_path.read_text(encoding="utf-8")
        assert ".exists()" in source
        assert "bundled_path" in source

    # ── 4f. Config path structure ─────────────────────────────────────────

    def test_coref_spacy_model_under_bundled_models(self):
        """SPACY_EN_CORE_WEB_SM_PATH is under BUNDLED_MODELS_DIR."""
        from src.config import BUNDLED_MODELS_DIR, SPACY_EN_CORE_WEB_SM_PATH

        rel = SPACY_EN_CORE_WEB_SM_PATH.relative_to(BUNDLED_MODELS_DIR)
        assert str(rel), "Path should be relative to BUNDLED_MODELS_DIR"
