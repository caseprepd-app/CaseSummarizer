"""
Tests for Word export content verification.

Verifies that exported .docx files contain the expected content
by reading them back with python-docx.
"""

from unittest.mock import MagicMock, patch

from docx import Document

# ============================================================================
# WordDocumentBuilder content verification
# ============================================================================


class TestWordBuilderContent:
    """Verify WordDocumentBuilder produces correct .docx content."""

    def _builder(self):
        from src.core.export.word_builder import WordDocumentBuilder

        return WordDocumentBuilder()

    def test_heading_appears_in_doc(self, tmp_path):
        builder = self._builder()
        builder.add_heading("Test Heading", level=1)
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Test Heading" in headings

    def test_paragraph_text(self, tmp_path):
        builder = self._builder()
        builder.add_paragraph("Hello world")
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "Hello world" in texts

    def test_bold_paragraph(self, tmp_path):
        builder = self._builder()
        builder.add_paragraph("Bold text", bold=True)
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        for p in doc.paragraphs:
            if p.text == "Bold text":
                assert p.runs[0].bold is True
                break
        else:
            raise AssertionError("Bold paragraph not found")

    def test_table_headers_and_rows(self, tmp_path):
        builder = self._builder()
        builder.add_table(
            headers=["Term", "Score"],
            rows=[["plaintiff", "80"], ["defendant", "75"]],
        )
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Header row
        assert table.rows[0].cells[0].text == "Term"
        assert table.rows[0].cells[1].text == "Score"
        # Data rows
        assert table.rows[1].cells[0].text == "plaintiff"
        assert table.rows[2].cells[0].text == "defendant"

    def test_empty_headers_no_table(self, tmp_path):
        builder = self._builder()
        builder.add_table(headers=[], rows=[])
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        assert len(doc.tables) == 0

    def test_separator_appears(self, tmp_path):
        builder = self._builder()
        builder.add_separator()
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        separator_found = any("─" in p.text for p in doc.paragraphs)
        assert separator_found

    def test_styled_paragraph_with_color(self, tmp_path):
        from src.core.export.base import TextSpan

        builder = self._builder()
        spans = [
            TextSpan(text="green text", color=(0, 128, 0)),
            TextSpan(text=" red text", color=(255, 0, 0), bold=True),
        ]
        builder.add_styled_paragraph(spans)
        out = tmp_path / "test.docx"
        builder.save(str(out))

        doc = Document(str(out))
        # Find paragraph with colored runs
        for p in doc.paragraphs:
            if "green text" in p.text:
                assert len(p.runs) == 2
                assert p.runs[0].font.color.rgb is not None
                assert p.runs[1].bold is True
                break
        else:
            raise AssertionError("Styled paragraph not found")


# ============================================================================
# ExportService vocabulary to Word end-to-end
# ============================================================================


class TestExportServiceVocabWord:
    """ExportService.export_vocabulary_to_word produces readable .docx."""

    def test_vocab_export_roundtrip(self, tmp_path):
        from src.services.export_service import ExportService

        svc = ExportService()
        vocab = [
            {
                "Term": "myocardial infarction",
                "Score": "85",
                "Is Person": "No",
                "Found By": "NER",
                "Frequency": "3",
                "Category": "Medical",
            },
            {
                "Term": "John Smith",
                "Score": "90",
                "Is Person": "Yes",
                "Found By": "NER",
                "Frequency": "5",
                "Category": "Person",
            },
        ]
        out = tmp_path / "vocab.docx"

        with patch("src.services.export_service._auto_open_file"):
            success, _ = svc.export_vocabulary_to_word(vocab, str(out))

        assert success is True
        assert out.exists()

        # Read back and verify content
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += " " + cell.text

        combined = all_text + table_text
        assert "myocardial infarction" in combined
        assert "John Smith" in combined


# ============================================================================
# ExportService Q&A to Word end-to-end
# ============================================================================


class TestExportServiceSemanticWord:
    """ExportService.export_semantic_to_word produces readable .docx."""

    def test_semantic_export_roundtrip(self, tmp_path):
        from src.services.export_service import ExportService

        svc = ExportService()
        result = MagicMock()
        result.question = "What injuries occurred?"
        result.quick_answer = "The plaintiff suffered a broken arm."
        result.citation = "Page 12, lines 5-10"
        result.source_summary = "complaint.pdf"
        result.verification = None

        out = tmp_path / "qa.docx"
        with patch("src.services.export_service._auto_open_file"):
            success, _ = svc.export_semantic_to_word([result], str(out))

        assert success is True
        assert out.exists()

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "What injuries occurred?" in all_text
        assert "broken arm" in all_text


# ============================================================================
# ExportService combined to Word end-to-end
# ============================================================================


class TestExportServiceCombinedWord:
    """ExportService.export_combined_to_word produces readable .docx."""

    def test_combined_export_has_both_sections(self, tmp_path):
        from src.services.export_service import ExportService

        svc = ExportService()
        vocab = [
            {
                "Term": "Aspirin",
                "Score": "70",
                "Is Person": "No",
                "Found By": "RAKE",
                "Frequency": "2",
            }
        ]
        qa = MagicMock()
        qa.question = "What medication?"
        qa.quick_answer = "Aspirin was prescribed."
        qa.citation = "p3"
        qa.source_summary = "records.pdf"
        qa.verification = None

        out = tmp_path / "combined.docx"
        with patch("src.services.export_service._auto_open_file"):
            success, _ = svc.export_combined_to_word(vocab, [qa], str(out))

        assert success is True

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += " " + cell.text

        combined = all_text + table_text
        assert "Aspirin" in combined
        assert "What medication?" in combined
