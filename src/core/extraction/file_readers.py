"""
File Readers for Non-PDF Document Formats.

Provides text extraction for TXT, RTF, DOCX, and image files (PNG/JPG).
Each reader returns an ExtractionResult with text, confidence, and status.

Supported formats:
    - TXT: Direct file read (UTF-8)
    - RTF: striprtf library conversion
    - DOCX: python-docx paragraph and table extraction
    - PNG/JPG: OCR via Tesseract (delegates to OCRProcessor)

Example usage:
    >>> from src.core.extraction.dictionary_utils import DictionaryTextValidator
    >>> from src.core.extraction.ocr_processor import OCRProcessor
    >>> readers = FileReaders(DictionaryTextValidator(), OCRProcessor(DictionaryTextValidator()))
    >>> result = readers.read_text_file(Path("document.txt"))
    >>> print(result['method'])  # 'direct_read'
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

from PIL import Image

logger = logging.getLogger(__name__)

from .dictionary_utils import DictionaryTextValidator

if TYPE_CHECKING:
    from .extraction_result import ExtractionResult


class FileReaders:
    """
    Reads text from non-PDF document formats.

    Provides a unified interface for extracting text from TXT, RTF, DOCX,
    and image files. Each method returns an ExtractionResult.

    Attributes:
        dictionary: DictionaryTextValidator for confidence calculation
        ocr_processor: Optional OCRProcessor for image files
    """

    def __init__(self, dictionary: DictionaryTextValidator, ocr_processor=None):
        """
        Initialize the file readers.

        Args:
            dictionary: DictionaryTextValidator instance for confidence calculation
            ocr_processor: Optional OCRProcessor for image files
        """
        self.dictionary = dictionary
        self.ocr_processor = ocr_processor

    def read_text_file(self, file_path: Path) -> ExtractionResult:
        """
        Read a plain text (.txt) file.

        Args:
            file_path: Path to the text file

        Returns:
            ExtractionResult with text, method='direct_read', confidence.

        Example:
            >>> readers = FileReaders(DictionaryTextValidator())
            >>> result = readers.read_text_file(Path("notes.txt"))
            >>> print(f"Read {len(result['text'])} characters")
        """
        from .extraction_result import ExtractionResult

        logger.debug("Reading text file: %s", file_path.name)

        try:
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                text = f.read()

            confidence = self.dictionary.calculate_confidence(text)
            logger.debug("Text file dictionary confidence: %.1f%%", confidence)

            return ExtractionResult.success(
                text,
                "direct_read",
                confidence,
                page_count=1,
            )

        except Exception as e:
            return ExtractionResult.error(
                f"Failed to read text file: {e!s}",
                page_count=0,
            )

    def read_rtf_file(self, file_path: Path) -> ExtractionResult:
        """
        Read a Rich Text Format (.rtf) file.

        Uses striprtf library to convert RTF to plain text.

        Args:
            file_path: Path to the RTF file

        Returns:
            ExtractionResult with text, method='rtf_extraction', confidence.

        Example:
            >>> readers = FileReaders(DictionaryTextValidator())
            >>> result = readers.read_rtf_file(Path("document.rtf"))
        """
        from .extraction_result import ExtractionResult

        logger.debug("Reading RTF file: %s", file_path.name)

        try:
            from striprtf.striprtf import rtf_to_text

            with open(file_path, encoding="utf-8", errors="ignore") as f:
                rtf_content = f.read()

            text = rtf_to_text(rtf_content)
            logger.debug("Extracted %d characters from RTF", len(text))

            confidence = self.dictionary.calculate_confidence(text)
            logger.debug("RTF dictionary confidence: %.1f%%", confidence)

            return ExtractionResult.success(
                text,
                "rtf_extraction",
                confidence,
                page_count=1,
            )

        except Exception as e:
            return ExtractionResult.error(
                f"Failed to read RTF file: {e!s}",
                page_count=0,
            )

    def read_docx_file(self, file_path: Path) -> ExtractionResult:
        """
        Read a Word document (.docx) file.

        Extracts text from paragraphs and tables using python-docx.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ExtractionResult with text, method='docx_extraction', confidence.

        Example:
            >>> readers = FileReaders(DictionaryTextValidator())
            >>> result = readers.read_docx_file(Path("report.docx"))
            >>> print(f"Pages: ~{result['page_count']}")
        """
        from .extraction_result import ExtractionResult

        logger.debug("Reading Word document: %s", file_path.name)

        try:
            from docx import Document

            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += "\n" + " | ".join(row_text)

            if not text.strip():
                return ExtractionResult.error(
                    "Word document contains no readable text.",
                    page_count=0,
                )

            confidence = self.dictionary.calculate_confidence(text)
            logger.debug("DOCX dictionary confidence: %.1f%%", confidence)

            return ExtractionResult.success(
                text,
                "docx_extraction",
                confidence,
                page_count=len(doc.sections) or 1,
            )

        except Exception as e:
            return ExtractionResult.error(
                f"Failed to read Word document: {e!s}",
                page_count=0,
            )

    def read_image_file(self, file_path: Path) -> ExtractionResult:
        """
        Read an image file (.png, .jpg, .jpeg) using OCR.

        Delegates to OCRProcessor for Tesseract-based text extraction.

        Args:
            file_path: Path to the image file

        Returns:
            ExtractionResult with text, method='image_ocr', confidence.

        Example:
            >>> readers = FileReaders(DictionaryTextValidator(), OCRProcessor(DictionaryTextValidator()))
            >>> result = readers.read_image_file(Path("scan.png"))
        """
        from .extraction_result import ExtractionResult

        logger.debug("Reading image file: %s", file_path.name)

        if self.ocr_processor is None:
            return ExtractionResult.error(
                "OCR processor not available for image files.",
                page_count=1,
            )

        try:
            with Image.open(file_path) as img:
                result = self.ocr_processor.process_image(img)

            # Add page_count to result
            result.page_count = 1

            return result

        except Exception as e:
            return ExtractionResult.error(
                f"Failed to process image: {e!s}",
                page_count=1,
            )
