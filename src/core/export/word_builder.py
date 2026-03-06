"""
Word Document Builder

Implements DocumentBuilder using python-docx for .docx export.
"""

import logging

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from src.core.export.base import DocumentBuilder, TextSpan


class WordDocumentBuilder(DocumentBuilder):
    """
    Word document builder using python-docx.

    Creates .docx files with proper formatting for vocabulary
    and Q&A exports including colored verification spans.
    """

    def __init__(self, title: str = "CasePrepd Export"):
        """
        Initialize a new Word document.

        Args:
            title: Document title (added as first heading)
        """
        self.doc = Document()
        self._set_document_styles()

    def _set_document_styles(self) -> None:
        """Configure default document styles."""
        # Set default font
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading to the document."""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> None:
        """Add a simple text paragraph."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic

    def add_styled_paragraph(self, spans: list[TextSpan]) -> None:
        """Add a paragraph with mixed styling (verification colors)."""
        para = self.doc.add_paragraph()

        for span in spans:
            run = para.add_run(span.text)

            if span.bold:
                run.bold = True
            if span.italic:
                run.italic = True
            if span.strikethrough:
                run.font.strike = True
            if span.color:
                run.font.color.rgb = RGBColor(*span.color)

    def add_table(self, headers: list[str], rows: list[list[str]]) -> None:
        """Add a table to the document."""
        if not headers:
            return

        # Create table with header row
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell_text)

    def add_separator(self) -> None:
        """Add a visual separator."""
        # Add horizontal line via paragraph border
        para = self.doc.add_paragraph()
        para.add_run("─" * 60)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, path: str) -> None:
        """Save the document to file."""
        try:
            self.doc.save(path)
        except OSError as e:
            raise OSError(
                f"Could not save document to '{path}'. "
                f"The file may be open in another application. "
                f"Please close it and try again."
            ) from e
