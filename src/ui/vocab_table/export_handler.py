"""
Vocabulary Export Handler Mixin.

Session 82: Extracted from dynamic_output.py for modularity.

Contains all export methods for vocabulary data:
- CSV export (with visible columns support)
- TXT export
- Word (DOCX) export
- PDF export
- HTML export
"""

import csv
import io
from datetime import datetime
from tkinter import filedialog, messagebox

from src.logging_config import debug_log
from src.ui.vocab_table.column_config import ALL_EXPORT_COLUMNS, DISPLAY_TO_DATA_COLUMN


class VocabExportMixin:
    """
    Mixin class providing vocabulary export functionality.

    Methods in this mixin assume the parent class has:
    - self._vocab_csv_data: List of vocabulary dictionaries
    - self._get_visible_columns(): Returns list of visible column names
    - self._get_sorted_vocab_data(): Returns sorted vocabulary data
    """

    def _build_vocab_csv(
        self, vocab_data: list[dict], visible_columns: list[str] | None = None
    ) -> str:
        """
        Build CSV string from vocabulary data.

        Args:
            vocab_data: List of vocabulary term dictionaries
            visible_columns: If provided, only include these columns (in order).
                           If None, includes all export columns.

        Returns:
            CSV string with headers and data rows
        """
        output = io.StringIO()

        # Session 80: Export only visible columns, preserving GUI column names
        columns = visible_columns or list(ALL_EXPORT_COLUMNS)

        writer = csv.DictWriter(
            output,
            fieldnames=columns,
            extrasaction="ignore",
            quoting=csv.QUOTE_MINIMAL,
        )
        writer.writeheader()

        for term in vocab_data:
            # Create row with mapped column names
            row = {}
            for col in columns:
                # Map display name to data field name if needed
                data_key = DISPLAY_TO_DATA_COLUMN.get(col, col)
                row[col] = term.get(data_key, term.get(col, ""))
            writer.writerow(row)

        return output.getvalue()

    def _export_vocab_csv(self):
        """
        Export visible vocabulary data to CSV file.

        Session 80: Respects current column visibility settings.
        Prompts user to choose filename and exports CSV.
        """
        if not hasattr(self, "_vocab_csv_data") or not self._vocab_csv_data:
            messagebox.showwarning("No Data", "No vocabulary data to export.")
            return

        # Get visible columns in display order
        visible_columns = self._get_visible_columns()

        # Use sorted data if available
        sorted_data = self._get_sorted_vocab_data()

        # Build CSV content
        csv_content = self._build_vocab_csv(sorted_data, visible_columns)

        # Prompt for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            initialfile=f"vocabulary_{timestamp}.csv",
            title="Export Vocabulary to CSV",
        )

        if filepath:
            with open(filepath, "w", encoding="utf-8", newline="") as f:
                f.write(csv_content)
            debug_log(f"[VocabExport] Exported {len(sorted_data)} terms to {filepath}")

    def _export_vocab_txt(self):
        """
        Export vocabulary to plain text file.

        Format: Tab-separated values with headers.
        """
        if not hasattr(self, "_vocab_csv_data") or not self._vocab_csv_data:
            messagebox.showwarning("No Data", "No vocabulary data to export.")
            return

        visible_columns = self._get_visible_columns()
        sorted_data = self._get_sorted_vocab_data()

        # Build TXT content
        lines = ["\t".join(visible_columns)]
        for term in sorted_data:
            row_values = []
            for col in visible_columns:
                data_key = DISPLAY_TO_DATA_COLUMN.get(col, col)
                value = term.get(data_key, term.get(col, ""))
                row_values.append(str(value))
            lines.append("\t".join(row_values))

        txt_content = "\n".join(lines)

        # Prompt for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            initialfile=f"vocabulary_{timestamp}.txt",
            title="Export Vocabulary to TXT",
        )

        if filepath:
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(txt_content)
            debug_log(f"[VocabExport] Exported {len(sorted_data)} terms to {filepath}")

    def _export_vocab_word(self):
        """
        Export vocabulary to Word document.

        Creates a formatted table with headers and data.
        """
        if not hasattr(self, "_vocab_csv_data") or not self._vocab_csv_data:
            messagebox.showwarning("No Data", "No vocabulary data to export.")
            return

        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.shared import Inches, Pt  # noqa: F401
        except ImportError:
            messagebox.showerror(
                "Missing Library",
                "python-docx is required for Word export.\n\nInstall with: pip install python-docx",
            )
            return

        visible_columns = self._get_visible_columns()
        sorted_data = self._get_sorted_vocab_data()

        # Create document
        doc = Document()
        doc.add_heading("Vocabulary Export", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total terms: {len(sorted_data)}")

        # Create table
        table = doc.add_table(rows=1, cols=len(visible_columns))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(visible_columns):
            header_cells[i].text = col
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Data rows
        for term in sorted_data:
            row_cells = table.add_row().cells
            for i, col in enumerate(visible_columns):
                data_key = DISPLAY_TO_DATA_COLUMN.get(col, col)
                value = term.get(data_key, term.get(col, ""))
                row_cells[i].text = str(value)

        # Prompt for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=f"vocabulary_{timestamp}.docx",
            title="Export Vocabulary to Word",
        )

        if filepath:
            doc.save(filepath)
            debug_log(f"[VocabExport] Exported {len(sorted_data)} terms to {filepath}")

    def _export_vocab_pdf(self):
        """
        Export vocabulary to PDF document.

        Creates a formatted table with headers and data.
        """
        if not hasattr(self, "_vocab_csv_data") or not self._vocab_csv_data:
            messagebox.showwarning("No Data", "No vocabulary data to export.")
            return

        try:
            from fpdf import FPDF
        except ImportError:
            messagebox.showerror(
                "Missing Library",
                "fpdf2 is required for PDF export.\n\nInstall with: pip install fpdf2",
            )
            return

        visible_columns = self._get_visible_columns()
        sorted_data = self._get_sorted_vocab_data()

        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Vocabulary Export", ln=True, align="C")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 8, f"Total terms: {len(sorted_data)}", ln=True)
        pdf.ln(5)

        # Calculate column widths (fit to page)
        page_width = pdf.w - 20  # margins
        col_count = len(visible_columns)
        col_width = page_width / col_count

        # Header row
        pdf.set_font("Helvetica", "B", 8)
        for col in visible_columns:
            pdf.cell(col_width, 8, col[:12], border=1, align="C")
        pdf.ln()

        # Data rows
        pdf.set_font("Helvetica", "", 7)
        for term in sorted_data:
            for col in visible_columns:
                data_key = DISPLAY_TO_DATA_COLUMN.get(col, col)
                value = str(term.get(data_key, term.get(col, "")))
                # Truncate for PDF cell
                if len(value) > 15:
                    value = value[:12] + "..."
                pdf.cell(col_width, 6, value, border=1, align="C")
            pdf.ln()

        # Prompt for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            initialfile=f"vocabulary_{timestamp}.pdf",
            title="Export Vocabulary to PDF",
        )

        if filepath:
            pdf.output(filepath)
            debug_log(f"[VocabExport] Exported {len(sorted_data)} terms to {filepath}")

    def _export_vocab_html(self):
        """
        Export vocabulary to HTML file.

        Session 80: Uses shared HTML builder with visible columns.
        Session 83: Updated to use ExportService.
        Creates a styled HTML table with sortable columns.
        """
        if not hasattr(self, "_vocab_csv_data") or not self._vocab_csv_data:
            messagebox.showwarning("No Data", "No vocabulary data to export.")
            return

        from src.services import ExportService

        visible_columns = self._get_visible_columns()
        sorted_data = self._get_sorted_vocab_data()

        # Prompt for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".html",
            filetypes=[("HTML files", "*.html"), ("All files", "*.*")],
            initialfile=f"vocabulary_{timestamp}.html",
            title="Export Vocabulary to HTML",
        )

        if filepath:
            export_service = ExportService()
            html_content = export_service.get_vocabulary_html_content(sorted_data, visible_columns)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(html_content)
            debug_log(f"[VocabExport] Exported {len(sorted_data)} terms to {filepath}")

    def _get_visible_columns(self) -> list[str]:
        """
        Get list of currently visible columns in display order.

        Override in parent class to provide actual visibility state.

        Returns:
            List of visible column names
        """
        # Default implementation - subclass should override
        from src.ui.vocab_table.column_config import COLUMN_ORDER

        return COLUMN_ORDER

    def _get_sorted_vocab_data(self) -> list[dict]:
        """
        Get vocabulary data in current sort order.

        Override in parent class to provide sorted data.

        Returns:
            List of vocabulary dictionaries
        """
        # Default implementation - subclass should override
        return getattr(self, "_vocab_csv_data", [])
